from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "orbbt_website_development_brief.md"
OUTPUT = ROOT / "docs" / "orbbt_website_development_brief.docx"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_paragraph(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_text(paragraph, text):
    url_match = re.fullmatch(r"https?://\S+", text.strip())
    if url_match:
        add_hyperlink(paragraph, text.strip(), text.strip())
        return

    stripped = text.replace("`", "")
    if ":" in stripped and len(stripped.split(":", 1)[0]) <= 42:
        label, rest = stripped.split(":", 1)
        run = paragraph.add_run(label + ":")
        set_run_font(run, bold=True)
        if rest:
            run = paragraph.add_run(rest)
            set_run_font(run)
    else:
        run = paragraph.add_run(stripped)
        set_run_font(run)


def add_numbering_style(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    add_numbering_style(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("orbbt. Website Development Brief")
    set_run_font(run, size=9, color="6D6A64")

    previous_blank = False
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line:
            previous_blank = True
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            style_paragraph(paragraph, before=0, after=8, line=1.0)
            run = paragraph.add_run(line[2:])
            set_run_font(run, size=24, color="111111", bold=True)
            previous_blank = False
            continue

        if line.startswith("## "):
            paragraph = doc.add_paragraph()
            style_paragraph(paragraph, before=14, after=6, line=1.10)
            run = paragraph.add_run(line[3:])
            set_run_font(run, size=16, color="2E74B5", bold=True)
            previous_blank = False
            continue

        if line.startswith("### "):
            paragraph = doc.add_paragraph()
            style_paragraph(paragraph, before=10, after=5, line=1.10)
            run = paragraph.add_run(line[4:])
            set_run_font(run, size=13, color="2E74B5", bold=True)
            previous_blank = False
            continue

        if line.startswith("#### "):
            paragraph = doc.add_paragraph()
            style_paragraph(paragraph, before=8, after=4, line=1.10)
            run = paragraph.add_run(line[5:])
            set_run_font(run, size=12, color="1F4D78", bold=True)
            previous_blank = False
            continue

        stripped = line.lstrip()
        indent = len(line) - len(stripped)
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.25 + min(indent, 4) * 0.06)
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_text(paragraph, stripped[2:])
            previous_blank = False
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_text(paragraph, numbered.group(1))
            previous_blank = False
            continue

        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, before=0 if not previous_blank else 2, after=6, line=1.10)
        add_inline_text(paragraph, line)
        previous_blank = False

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
